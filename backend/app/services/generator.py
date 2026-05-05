import os
import re
import json
import time
import logging
from typing import Any, Optional

from groq import Groq
from tenacity import (
    retry,
    stop_after_attempt,
    wait_random_exponential,
    retry_if_exception_type,
)
from app.services.prompts import (
    PROMPT_FUNCTIONAL,
    PROMPT_JIRA_EXCEL,
    PROMPT_TECHNICAL,
    PROMPT_BUDGET_PLANNING,
)
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logger = logging.getLogger("mawaba.generator")
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# ---------------------------------------------------------------------------
# Configuration Groq
# ---------------------------------------------------------------------------
client = Groq(
    api_key=os.getenv("GROQ_API_KEY"),
    timeout=300.0,
)

MODEL_NAME = "llama-3.3-70b-versatile"
FALLBACK_MODEL = "llama-3.1-8b-instant"

# ---------------------------------------------------------------------------
# Dossiers de sortie
# ---------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
OUTPUT_ROOT = os.path.join(BASE_DIR, "outputs")
OUTPUT_WORD = os.path.join(OUTPUT_ROOT, "word")
OUTPUT_EXCEL = os.path.join(OUTPUT_ROOT, "excel")
OUTPUT_HTML = os.path.join(OUTPUT_ROOT, "html")

for _d in (OUTPUT_WORD, OUTPUT_EXCEL, OUTPUT_HTML):
    os.makedirs(_d, exist_ok=True)

# ---------------------------------------------------------------------------
# Couleurs & Styles Mawaba
# ---------------------------------------------------------------------------
COLOR_PRIMARY = RGBColor(30, 41, 59)       # Slate-800
COLOR_ACCENT = RGBColor(56, 189, 248)      # Sky-400
COLOR_MUTED = RGBColor(100, 116, 139)      # Slate-500
COLOR_WHITE = RGBColor(255, 255, 255)
COLOR_LIGHT_BG = RGBColor(241, 245, 249)   # Slate-100
HEX_PRIMARY = "1E293B"
HEX_ACCENT = "38BDF8"


# ═══════════════════════════════════════════════════════════════════════════
# UTILITAIRES
# ═══════════════════════════════════════════════════════════════════════════

def clean_json_response(text: str) -> str:
    """Extrait le JSON d'une réponse potentiellement encapsulée dans des backticks."""
    text = re.sub(r"^```(?:json)?\s*", "", text.strip(), flags=re.MULTILINE)
    text = re.sub(r"\s*```$", "", text.strip(), flags=re.MULTILINE)
    return text.strip()


def parse_llm_json(text: str) -> dict:
    """Parse robuste du JSON LLM avec fallback progressif."""
    cleaned = clean_json_response(text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    # Tentative sur le texte brut
    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        logger.error(f"Échec du parsing JSON : {e}\nDébut du texte : {cleaned[:300]}")
        raise ValueError(f"Impossible de parser la réponse JSON du LLM : {e}")


def get_robust(data: Any, keys_to_try: list[str], default: Any = "") -> Any:
    """Cherche une valeur dans un dictionnaire avec tolérance sur les clés."""
    if not isinstance(data, dict):
        return default
    # 1. Exact match
    for k in keys_to_try:
        if k in data:
            return data[k]
    # 2. Case-insensitive
    data_lower = {str(k).lower().strip(): v for k, v in data.items()}
    for k in keys_to_try:
        if k.lower().strip() in data_lower:
            return data_lower[k.lower().strip()]
    # 3. Substring match
    for k in keys_to_try:
        for dk, dv in data.items():
            if k.lower() in str(dk).lower():
                return dv
    return default


# ═══════════════════════════════════════════════════════════════════════════
# APPEL LLM
# ═══════════════════════════════════════════════════════════════════════════

@retry(
    stop=stop_after_attempt(4),
    wait=wait_random_exponential(multiplier=2, max=60),
    retry=retry_if_exception_type(Exception),
)
def call_llm(prompt: str, response_format: str = "text", model: Optional[str] = None) -> str:
    """Appel Groq avec retry automatique, rate-limit handling, et fallback modèle."""
    target_model = model or MODEL_NAME

    kwargs = {
        "model": target_model,
        "messages": [
            {
                "role": "system",
                "content": (
                    "Tu es un expert en gestion de projet IA chez Mawaba Technologies. "
                    "Tu produis des livrables professionnels de qualité cabinet de conseil. "
                    "Si on te demande du JSON, retourne UNIQUEMENT du JSON valide — "
                    "aucun texte, aucun commentaire, aucun markdown autour."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.15,
        "max_tokens": 8192,  # Doublé pour des réponses plus riches
    }
    if response_format == "json":
        kwargs["response_format"] = {"type": "json_object"}

    try:
        time.sleep(0.5)
        response = client.chat.completions.create(**kwargs)
        return response.choices[0].message.content
    except Exception as e:
        error_str = str(e).lower()
        if "rate_limit" in error_str or "429" in error_str:
            logger.warning(f"Rate limit sur {target_model}, pause 30s…")
            time.sleep(30)
        # Fallback sur modèle léger si le principal échoue
        if target_model == MODEL_NAME:
            logger.warning(f"Tentative fallback sur {FALLBACK_MODEL}…")
            kwargs["model"] = FALLBACK_MODEL
            kwargs["max_tokens"] = 2000
            try:
                response = client.chat.completions.create(**kwargs)
                return response.choices[0].message.content
            except Exception as fallback_err:
                logger.error(f"Échec fallback : {fallback_err}")
        raise


# ═══════════════════════════════════════════════════════════════════════════
# FORMATAGE WORD — Utilitaires de rendu professionnel
# ═══════════════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, color_hex: str):
    """Applique une couleur de fond à une cellule Word."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _style_table_header(table, header_color: str = HEX_PRIMARY):
    """Stylise la première ligne d'un tableau en header professionnel."""
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, header_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = COLOR_WHITE
                run.font.size = Pt(10)
                run.font.name = "Arial"


def _add_zebra_striping(table, even_color: str = "F1F5F9"):
    """Ajoute un fond alterné aux lignes paires du tableau."""
    for i, row in enumerate(table.rows[1:], start=1):
        if i % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, even_color)


def setup_doc_style(doc: Document) -> Document:
    """Configure la police et les styles de base du document."""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(51, 51, 51)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Style des headings
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Arial"
        heading_style.font.color.rgb = COLOR_PRIMARY
        if level == 1:
            heading_style.font.size = Pt(16)
        elif level == 2:
            heading_style.font.size = Pt(13)
        else:
            heading_style.font.size = Pt(11)

    return doc


def _add_cover_page(doc: Document, title: str, subtitle: str, project_name: str):
    """Ajoute une page de garde professionnelle."""
    for _ in range(4):
        doc.add_paragraph()

    # Logo
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo.add_run("MAWABA TECHNOLOGIES")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.name = "Arial"

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = tagline.add_run("IA & Machine Learning - Développement Logiciels - DevOps & Automatisation - Cybersécurité")
    t_run.font.size = Pt(12)
    t_run.font.color.rgb = COLOR_MUTED
    t_run.font.name = "Arial"

    # Séparateur
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sep.add_run("━" * 30)
    s_run.font.color.rgb = COLOR_ACCENT
    s_run.font.size = Pt(14)

    doc.add_paragraph()

    # Titre du document
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = t.add_run(title)
    t_run.bold = True
    t_run.font.size = Pt(24)
    t_run.font.color.rgb = COLOR_PRIMARY
    t_run.font.name = "Arial"

    # Sous-titre
    if subtitle:
        st = doc.add_paragraph()
        st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st_run = st.add_run(subtitle)
        st_run.font.size = Pt(14)
        st_run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()

    # Nom du projet
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p.add_run(f"Projet : {project_name}")
    p_run.font.size = Pt(16)
    p_run.font.color.rgb = COLOR_ACCENT
    p_run.font.name = "Arial"

    # Date
    from datetime import date
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = d.add_run(date.today().strftime("%d/%m/%Y"))
    d_run.font.size = Pt(11)
    d_run.font.color.rgb = COLOR_MUTED

    doc.add_page_break()


def _add_footer(doc: Document, project_name: str):
    """Ajoute un pied de page professionnel avec numérotation."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(f"Mawaba Technologies — {project_name} — Confidentiel  |  Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED
        run.font.name = "Arial"

        # Numéro de page automatique
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")

        run2 = p.add_run()
        run2.font.size = Pt(8)
        run2.font.color.rgb = COLOR_MUTED
        run2._r.append(fld_char1)
        run2._r.append(instr)
        run2._r.append(fld_char2)


def _format_value_to_paragraphs(doc: Document, value: Any, level: int = 0):
    """Convertit récursivement une valeur JSON en paragraphes Word formatés."""
    indent = Cm(level * 0.8)

    if isinstance(value, str):
        if value.strip():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = indent
            run = p.add_run(value)
            run.font.size = Pt(10)

    elif isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                _format_dict_as_block(doc, item, level)
            elif isinstance(item, str):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = indent
                run = p.runs[0] if p.runs else p.add_run(item)
                if not p.runs:
                    pass
                else:
                    run.font.size = Pt(10)
                # Re-set text if needed
                if p.text != item:
                    p.clear()
                    r = p.add_run(item)
                    r.font.size = Pt(10)
            else:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = indent
                r = p.add_run(str(item))
                r.font.size = Pt(10)

    elif isinstance(value, dict):
        _format_dict_as_block(doc, value, level)

    else:
        if value is not None:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = indent
            run = p.add_run(str(value))
            run.font.size = Pt(10)


def _format_dict_as_block(doc: Document, data: dict, level: int = 0):
    """Formate un dictionnaire en bloc structuré dans le document."""
    indent = Cm(level * 0.8)
    for key, val in data.items():
        # Titre de la clé
        clean_key = str(key).replace("_", " ").title()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(f"▸ {clean_key} : ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_PRIMARY

        if isinstance(val, (str, int, float)):
            val_run = p.add_run(str(val))
            val_run.font.size = Pt(10)
        elif isinstance(val, list) and all(isinstance(x, str) for x in val):
            val_run = p.add_run(", ".join(val))
            val_run.font.size = Pt(10)
        elif isinstance(val, list):
            for item in val:
                _format_value_to_paragraphs(doc, item, level + 1)
        elif isinstance(val, dict):
            _format_dict_as_block(doc, val, level + 1)
        elif val is None:
            val_run = p.add_run("À préciser")
            val_run.font.size = Pt(10)
            val_run.font.italic = True
            val_run.font.color.rgb = COLOR_MUTED


def _format_section_content(doc: Document, content: Any):
    """Formate intelligemment le contenu d'une section selon son type."""
    if isinstance(content, str):
        # Essai de parse JSON si c'est un dict stringifié
        try:
            parsed = json.loads(content.replace("'", '"'))
            _format_value_to_paragraphs(doc, parsed)
            return
        except (json.JSONDecodeError, ValueError):
            pass

        # Vérifier si c'est un repr() Python de dict/list
        if content.startswith("{") or content.startswith("["):
            try:
                import ast
                parsed = ast.literal_eval(content)
                _format_value_to_paragraphs(doc, parsed)
                return
            except (ValueError, SyntaxError):
                pass

        # Texte brut — ajouter tel quel
        for line in content.split("\n"):
            if line.strip():
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.font.size = Pt(10)

    elif isinstance(content, (dict, list)):
        _format_value_to_paragraphs(doc, content)
    else:
        p = doc.add_paragraph()
        run = p.add_run(str(content))
        run.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════════════
# CRÉATION DES DOCUMENTS
# ═══════════════════════════════════════════════════════════════════════════

def create_word_file(
    filename: str,
    title: str,
    content_sections: dict[str, Any],
    project_name: str = "Projet IA",
    subtitle: str = "",
) -> str:
    """Crée un document Word professionnel avec mise en forme riche."""
    doc = Document()
    setup_doc_style(doc)

    # Marges
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Page de garde
    _add_cover_page(doc, title, subtitle, project_name)

    # Table des matières (placeholder textuel)
    doc.add_heading("Table des Matières", level=1)
    for i, section_title in enumerate(content_sections.keys(), 1):
        clean_title = section_title.split(". ", 1)[-1] if ". " in section_title else section_title
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {clean_title}")
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_PRIMARY
    doc.add_page_break()

    # Contenu
    for counter, (section_title, content) in enumerate(content_sections.items(), 1):
        clean_title = section_title.split(". ", 1)[-1] if ". " in section_title else section_title
        numbered_title = f"{counter}. {clean_title}"

        h = doc.add_heading(numbered_title, level=1)
        for run in h.runs:
            run.font.color.rgb = COLOR_PRIMARY

        # Séparateur fin sous le titre
        sep = doc.add_paragraph()
        s_run = sep.add_run("━" * 50)
        s_run.font.color.rgb = COLOR_ACCENT
        s_run.font.size = Pt(6)

        _format_section_content(doc, content)
        doc.add_paragraph()  # Espacement

    # Pied de page
    _add_footer(doc, project_name)

    path = os.path.join(OUTPUT_WORD, filename)
    doc.save(path)
    logger.info(f"Document Word créé : {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════
# GÉNÉRATEURS DE DOCUMENTS
# ═══════════════════════════════════════════════════════════════════════════

def generate_functional_doc(project_info: dict) -> tuple[str, dict]:
    """Génère le Cahier des Charges Fonctionnel (CdCF)."""
    logger.info(f"Génération CdCF : {project_info['titre']}")
    prompt = PROMPT_FUNCTIONAL.format(**project_info)
    response = call_llm(prompt, response_format="json")
    data = parse_llm_json(response)

    # On passe directement les objets JSON, plus de str()
    sections = {
        "B. Contexte du Projet": get_robust(data, ["B_contexte_projet"], {}),
        "C. Objectifs Stratégiques": get_robust(data, ["C_objectifs"], {}),
        "D. Périmètre et Hypothèses": get_robust(data, ["D_perimetre"], {}),
        "E. Profils Utilisateurs": get_robust(data, ["E_utilisateurs"], {}),
        "F. Cas d'Usage Détaillés": get_robust(data, ["F_cas_usage"], []),
        "G. Exigences Fonctionnelles": get_robust(data, ["G_exigences_fonctionnelles"], {}),
        "H. Exigences Non Fonctionnelles": get_robust(data, ["H_exigences_non_fonctionnelles", "H_exigences_non_fictionnelles"], {}),
        "I. Critères d'Acceptation": get_robust(data, ["I_criteres_acceptation"], []),
        "J. KPIs et Mesure du Succès": get_robust(data, ["J_kpis_succes"], []),
    }

    path = create_word_file(
        "01_Cahier_Charges_Fonctionnel.docx",
        "Cahier des Charges Fonctionnel",
        sections,
        project_name=project_info["titre"],
        subtitle="Document de référence fonctionnel",
    )
    return path, data


def generate_technical_doc(project_info: dict, functional_data: dict) -> tuple[str, dict]:
    """Génère la Note d'Architecture Technique (DAT)."""
    logger.info(f"Génération DAT : {project_info['titre']}")
    prompt = PROMPT_TECHNICAL.format(
        titre=project_info["titre"],
        context_fonctionnel=json.dumps(get_robust(functional_data, ["B_contexte_projet"], {}), ensure_ascii=False),
        cas_usage=json.dumps(get_robust(functional_data, ["F_cas_usage"], []), ensure_ascii=False),
        exigences_non_fonctionnelles=json.dumps(get_robust(functional_data, ["H_exigences_non_fonctionnelles"], {}), ensure_ascii=False),
    )
    response = call_llm(prompt, response_format="json")
    data = parse_llm_json(response)

    sections = {
        "A. Présentation Générale": get_robust(data, ["A_presentation_generale"], {}),
        "B. Architecture Système": get_robust(data, ["B_architecture_systeme"], {}),
        "C. Architecture IA / RAG": get_robust(data, ["C_architecture_ia"], {}),
        "D. Sources de Données": get_robust(data, ["D_sources_de_donnees"], []),
        "E. Pipeline de Traitement": get_robust(data, ["E_pipeline_traitement"], {}),
        "F. Stack Technologique": get_robust(data, ["F_stack_technologique"], {}),
        "G. Sécurité Technique": get_robust(data, ["G_securite_technique"], {}),
        "H. Contraintes et Limites": get_robust(data, ["H_contraintes_techniques"], {}),
        "I. Stratégie de Déploiement": get_robust(data, ["I_deploiement"], {}),
        "J. Maintenance et Observabilité": get_robust(data, ["J_maintenance_observabilite"], {}),
    }

    path = create_word_file(
        "03_Document_Technique.docx",
        "Note d'Architecture Technique (DAT)",
        sections,
        project_name=project_info["titre"],
        subtitle="Architecture & Choix Techniques",
    )
    return path, data


def generate_budget_planning_doc(project_info: dict, functional_data: dict, technical_data: dict) -> str:
    """Génère le Budget & Planning Stratégique avec tableaux stylisés."""
    logger.info(f"Génération Budget/Planning : {project_info['titre']}")
    prompt = PROMPT_BUDGET_PLANNING.format(
        titre=project_info["titre"],
        echeancier=project_info["echeancier"],
        budget_max=project_info.get("budget_max", "Non défini"),
        context_fonctionnel=json.dumps(get_robust(functional_data, ["C_objectifs"], {}), ensure_ascii=False),
        stack_technique=json.dumps(get_robust(technical_data, ["F_stack_technologique"], {}), ensure_ascii=False),
    )
    response = call_llm(prompt, response_format="json")
    data = parse_llm_json(response)

    doc = Document()
    setup_doc_style(doc)

    # Marges
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Page de garde
    _add_cover_page(doc, "Budget & Planning Stratégique", "Estimation financière et planification", project_info["titre"])

    # --- Hypothèses budgétaires ---
    meta = get_robust(data, ["meta"], {})
    hypotheses = get_robust(meta, ["hypotheses_budgetaires"], [])
    if hypotheses:
        doc.add_heading("Hypothèses Budgétaires", level=1)
        for h in hypotheses:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            run = p.add_run(str(h))
            run.font.size = Pt(10)

    # --- 1. Budget RH ---
    doc.add_heading("1. Investissement Capital Humain", level=1)
    rh_data = get_robust(data, ["budget_rh"], {})
    desc_rh = get_robust(rh_data, ["description"], "")
    if desc_rh:
        p = doc.add_paragraph()
        run = p.add_run(desc_rh)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = COLOR_MUTED

    postes_rh = get_robust(rh_data, ["postes"], [])
    if postes_rh:
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["Poste", "Profil", "Jours", "TJM Min", "TJM Max", "Coût Max"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for item in postes_rh:
            cells = table.add_row().cells
            cells[0].text = str(item.get("poste", ""))
            cells[1].text = str(item.get("profil", ""))
            cells[2].text = str(item.get("jours_estimes", ""))
            cells[3].text = f"{item.get('tjm_min', 0):,} €".replace(",", " ")
            cells[4].text = f"{item.get('tjm_max', 0):,} €".replace(",", " ")
            cells[5].text = f"{item.get('cout_total_max', 0):,} €".replace(",", " ")

        _style_table_header(table)
        _add_zebra_striping(table)

    # Totaux RH
    total_min = get_robust(rh_data, ["total_rh_min"], 0)
    total_max = get_robust(rh_data, ["total_rh_max"], 0)
    p = doc.add_paragraph()
    run = p.add_run(f"💰 Total RH : {total_min:,} € — {total_max:,} €".replace(",", " "))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_PRIMARY

    # --- 2. Budget Technique ---
    doc.add_heading("2. Investissement Technique & Cloud (Annuel)", level=1)
    tech_budget = get_robust(data, ["budget_technique"], {})
    postes_tech = get_robust(tech_budget, ["postes"], [])
    if postes_tech:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["Item", "Détail", "Coût Mensuel Max", "Coût Annuel Max"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for item in postes_tech:
            cells = table.add_row().cells
            cells[0].text = str(item.get("item", ""))
            cells[1].text = str(item.get("detail", ""))
            cells[2].text = f"{item.get('cout_mensuel_max', 0):,} €".replace(",", " ")
            cells[3].text = f"{item.get('cout_annuel_max', 0):,} €".replace(",", " ")

        _style_table_header(table)
        _add_zebra_striping(table)

    # --- 3. Budget Total ---
    total_projet = get_robust(data, ["budget_total_projet"], {})
    if total_projet:
        doc.add_heading("3. Synthèse Budgétaire", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Poste"
        table.rows[0].cells[1].text = "Montant"

        rows_data = [
            ("Investissement initial (min)", f"{total_projet.get('investissement_initial_min', 0):,} €"),
            ("Investissement initial (max)", f"{total_projet.get('investissement_initial_max', 0):,} €"),
            ("Charge récurrente annuelle (min)", f"{total_projet.get('charge_recurrente_annuelle_min', 0):,} €"),
            ("Charge récurrente annuelle (max)", f"{total_projet.get('charge_recurrente_annuelle_max', 0):,} €"),
        ]
        for label, val in rows_data:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = val.replace(",", " ")

        _style_table_header(table, HEX_ACCENT)
        _add_zebra_striping(table)

        note = total_projet.get("note", "")
        if note:
            p = doc.add_paragraph()
            run = p.add_run(f"📝 {note}")
            run.font.size = Pt(10)
            run.font.italic = True

    # --- 4. Planning ---
    doc.add_heading("4. Chronogramme et Jalons Clés", level=1)
    phases = get_robust(data, ["planning_phases"], [])
    if phases:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["Phase", "Durée", "Objectif", "Livrable", "Risque"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for item in phases:
            cells = table.add_row().cells
            cells[0].text = str(item.get("phase", ""))
            cells[1].text = str(item.get("duree", ""))
            cells[2].text = str(item.get("objectif", ""))
            cells[3].text = str(item.get("livrable_principal", ""))
            cells[4].text = str(item.get("risque_phase", ""))

        _style_table_header(table)
        _add_zebra_striping(table)

    # --- 5. Jalons ---
    jalons = get_robust(data, ["jalons"], [])
    if jalons:
        doc.add_heading("5. Jalons Clés", level=1)
        for j in jalons:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            run = p.add_run(f"{j.get('id', '')} — {j.get('nom', '')} ({j.get('date_estimee', '')})")
            run.bold = True
            run.font.size = Pt(10)
            if j.get("critere_succes"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(1)
                r2 = p2.add_run(f"Critère : {j['critere_succes']}")
                r2.font.size = Pt(9)
                r2.font.color.rgb = COLOR_MUTED

    # --- 6. Risques ---
    risques = get_robust(data, ["registre_risques_budget"], [])
    if risques:
        doc.add_heading("6. Registre des Risques Budgétaires", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["Risque", "Probabilité", "Impact", "Mitigation"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for r in risques:
            cells = table.add_row().cells
            cells[0].text = str(r.get("risque", ""))
            cells[1].text = str(r.get("probabilite", ""))
            cells[2].text = str(r.get("impact", ""))
            cells[3].text = str(r.get("mitigation", ""))

        _style_table_header(table)
        _add_zebra_striping(table)

    _add_footer(doc, project_info["titre"])

    path = os.path.join(OUTPUT_WORD, "02_Budget_Planning.docx")
    doc.save(path)
    logger.info(f"Budget/Planning créé : {path}")
    return path


def generate_jira_excel(functional_data: dict, project_title: str) -> str:
    """Génère le Backlog Jira (Epics, Stories) en Excel stylisé."""
    logger.info("Génération Backlog Jira Excel…")
    import pandas as pd
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    prompt = PROMPT_JIRA_EXCEL.format(
        document_fonctionnel_json=json.dumps(functional_data, ensure_ascii=False)
    )
    response = call_llm(prompt, response_format="json")
    data = parse_llm_json(response)

    epics = get_robust(data, ["epics"], [])
    header = ["Type", "Issue Key", "Summary", "Description", "Priority", "Status", "Labels", "Story Points", "Sprint"]
    synthese_rows = [header]
    import_jira_rows = []
    epics_sheets = {}

    for epic in epics:
        epic_name = epic.get("summary", "Unknown Epic")
        epic_id = epic.get("id", "EPIC-X")
        epic_desc = epic.get("description", "")
        epic_prio = epic.get("priority", "High")
        epic_labels = ", ".join(epic.get("labels", ["IA"]))

        epic_rows = [header]
        epic_row = ["Epic", epic_id, epic_name, epic_desc, epic_prio, "To Do", epic_labels, "", ""]

        synthese_rows.append(epic_row)
        import_jira_rows.append({
            "Issue Key": epic_id,
            "Summary": epic_name,
            "Description": epic_desc,
            "Issue Type": "Epic",
            "Priority": epic_prio,
            "Status": "To Do",
            "Epic Link": "",
            "Labels": epic_labels,
        })
        epic_rows.append(epic_row)

        for us in epic.get("user_stories", []):
            us_row = [
                "Story",
                us.get("id", "US-XX"),
                us.get("summary", ""),
                us.get("description", ""),
                us.get("priority", "Medium"),
                "To Do",
                ", ".join(us.get("labels", ["US"])),
                us.get("story_points", ""),
                us.get("sprint_suggere", ""),
            ]
            synthese_rows.append(us_row)
            import_jira_rows.append({
                "Issue Key": us.get("id", ""),
                "Summary": us.get("summary", ""),
                "Description": us.get("description", ""),
                "Issue Type": "Story",
                "Priority": us.get("priority", "Medium"),
                "Status": "To Do",
                "Epic Link": epic_id,
                "Labels": ", ".join(us.get("labels", [])),
            })
            epic_rows.append(us_row)

        # Tâches techniques
        for task in epic.get("tasks_techniques", []):
            task_row = [
                "Task",
                task.get("id", "TASK-XX"),
                task.get("summary", ""),
                task.get("description", ""),
                task.get("priority", "Medium"),
                "To Do",
                ", ".join(task.get("labels", ["Tech"])),
                task.get("story_points", ""),
                "",
            ]
            synthese_rows.append(task_row)
            epic_rows.append(task_row)

        epics_sheets[epic_id] = epic_rows

    # Sprint Planning
    sprint_planning = get_robust(data, ["sprint_planning_suggere", "sprint_planning", "sprints"], [])

    path = os.path.join(OUTPUT_EXCEL, "04_User_Stories_Backlog.xlsx")

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        # 1. Synthèse
        pd.DataFrame(synthese_rows).to_excel(writer, sheet_name="Synthèse", index=False, header=False)

        # 2. Par Epic
        for eid, rows in epics_sheets.items():
            sheet_name = f"Module {eid}"[:31]
            pd.DataFrame(rows).to_excel(writer, sheet_name=sheet_name, index=False, header=False)

        # 3. Import Jira
        pd.DataFrame(import_jira_rows).to_excel(writer, sheet_name="Import Jira", index=False)

        # 4. Sprint Planning
        if sprint_planning:
            sp_rows = [["Sprint", "Objectif", "Stories", "Points", "Durée"]]
            for sp in sprint_planning:
                sp_rows.append([
                    sp.get("sprint", ""),
                    sp.get("objectif", ""),
                    ", ".join(sp.get("user_stories_ids", [])),
                    sp.get("story_points_total", ""),
                    sp.get("duree", ""),
                ])
            pd.DataFrame(sp_rows).to_excel(writer, sheet_name="Sprint Planning", index=False, header=False)

        # --- Stylisation ---
        workbook = writer.book
        header_fill = PatternFill(start_color=HEX_PRIMARY, end_color=HEX_PRIMARY, fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, name="Arial", size=10)
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style="thin", color="D1D5DB"),
            right=Side(style="thin", color="D1D5DB"),
            top=Side(style="thin", color="D1D5DB"),
            bottom=Side(style="thin", color="D1D5DB"),
        )
        even_fill = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")

        for sheet in workbook.worksheets:
            # Header
            for cell in sheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = thin_border

            # Data rows
            for row_idx, row in enumerate(sheet.iter_rows(min_row=2), start=2):
                for cell in row:
                    cell.alignment = Alignment(vertical="center", wrap_text=True)
                    cell.border = thin_border
                    cell.font = Font(name="Arial", size=9)
                    if row_idx % 2 == 0:
                        cell.fill = even_fill

            # Filtres et gel
            sheet.auto_filter.ref = sheet.dimensions
            sheet.freeze_panes = "A2"

            # Largeurs de colonnes
            for col in sheet.columns:
                max_length = 0
                col_letter = col[0].column_letter
                for cell in col:
                    try:
                        if cell.value and len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except (TypeError, AttributeError):
                        pass
                sheet.column_dimensions[col_letter].width = min(max_length + 3, 55)

    logger.info(f"Backlog Excel créé : {path}")
    return path


def generate_html_diagram(project_info: dict) -> str:
    """Génère le diagramme d'architecture interactif en HTML/Mermaid."""
    logger.info("Génération diagramme Mermaid…")
    prompt = f"""Génère uniquement le code Mermaid JS (type graph TD) pour un schéma d'architecture du projet : {project_info['titre']}.
RÈGLES CRITIQUES :
- Utilise des guillemets doubles pour TOUS les textes (ex: A["Texte ici"]).
- Pas de caractères spéciaux comme (), [], {{}} dans les textes des nœuds.
- Pas de parenthèses dans les labels — utilise uniquement des guillemets doubles.
- Ne mets PAS de bloc ```mermaid.
- Retourne juste le code brut sans aucun texte avant ou après.
- Minimum 8 nœuds avec des relations claires.
- Utilise des sous-graphes (subgraph) pour organiser par couche (Frontend, Backend, IA, Data)."""

    try:
        mermaid_code = call_llm(prompt, model=FALLBACK_MODEL)
        mermaid_code = mermaid_code.replace("```mermaid", "").replace("```", "").strip()
        if not mermaid_code.lower().startswith("graph"):
            mermaid_code = "graph TD\n" + mermaid_code
    except Exception as e:
        logger.error(f"Erreur Mermaid : {e}")
        mermaid_code = (
            'graph TD\n'
            '  A["Client Web"] --> B["API Gateway"]\n'
            '  B --> C["Backend Service"]\n'
            '  C --> D["Base de données"]\n'
            '  C --> E["Service IA"]\n'
            '  E --> F["LLM API"]\n'
            '  E --> G["Vector Store"]'
        )

    html_content = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Architecture — {project_info['titre']}</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <script type="module">
        import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
        mermaid.initialize({{
            startOnLoad: true,
            theme: 'base',
            themeVariables: {{
                'primaryColor': '#38bdf8',
                'primaryTextColor': '#0f172a',
                'primaryBorderColor': '#1e293b',
                'lineColor': '#64748b',
                'secondaryColor': '#e2e8f0',
                'tertiaryColor': '#f8fafc',
                'fontSize': '14px',
                'fontFamily': 'Inter, sans-serif'
            }},
            securityLevel: 'loose'
        }});
    </script>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Inter', sans-serif;
            background: linear-gradient(135deg, #0f172a 0%, #1e293b 50%, #0f172a 100%);
            color: #f8fafc;
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 40px 20px;
        }}
        .container {{
            width: 100%;
            max-width: 1200px;
        }}
        .card {{
            background: rgba(30, 41, 59, 0.8);
            backdrop-filter: blur(20px);
            border: 1px solid rgba(56, 189, 248, 0.15);
            border-radius: 24px;
            padding: 48px;
            box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.5),
                        0 0 80px -20px rgba(56, 189, 248, 0.1);
        }}
        .header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 32px;
            padding-bottom: 24px;
            border-bottom: 1px solid rgba(56, 189, 248, 0.2);
        }}
        .header-left h1 {{
            font-size: 1.5rem;
            font-weight: 600;
            color: #f8fafc;
            letter-spacing: -0.025em;
        }}
        .header-left p {{
            color: #94a3b8;
            font-size: 0.875rem;
            margin-top: 4px;
        }}
        .badge {{
            background: linear-gradient(135deg, #38bdf8, #0ea5e9);
            color: #0f172a;
            padding: 6px 16px;
            border-radius: 9999px;
            font-size: 11px;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }}
        .mermaid-wrapper {{
            background: rgba(255, 255, 255, 0.97);
            padding: 40px;
            border-radius: 16px;
            margin-top: 16px;
            box-shadow: inset 0 2px 4px 0 rgba(0, 0, 0, 0.04);
        }}
        .mermaid {{
            text-align: center;
        }}
        .footer {{
            margin-top: 32px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            font-size: 12px;
            color: #64748b;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="card">
            <div class="header">
                <div class="header-left">
                    <h1>🏗️ Architecture Technique</h1>
                    <p>{project_info['titre']}</p>
                </div>
                <span class="badge">Mawaba AI</span>
            </div>
            <div class="mermaid-wrapper">
                <div class="mermaid">
{mermaid_code}
                </div>
            </div>
            <div class="footer">
                <span>Généré automatiquement par Mawaba AI</span>
                <span>Document Confidentiel — © 2025</span>
            </div>
        </div>
    </div>
</body>
</html>"""

    path = os.path.join(OUTPUT_HTML, "05_Architecture.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html_content)
    logger.info(f"Diagramme HTML créé : {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════
# ORCHESTRATEUR PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════

def generate_all_deliverables(project_info: dict) -> dict[str, str]:
    """Génère tous les livrables du projet de manière séquentielle chaînée."""
    logger.info(f"{'='*60}")
    logger.info(f"DÉMARRAGE GÉNÉRATION COMPLÈTE : {project_info['titre']}")
    logger.info(f"{'='*60}")

    # Nettoyage des anciens fichiers
    for folder in (OUTPUT_WORD, OUTPUT_EXCEL, OUTPUT_HTML):
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except OSError as e:
                logger.warning(f"Impossible de supprimer {file_path}: {e}")

    start = time.time()

    # Étape 1 : CdCF (source de vérité)
    logger.info("━━━ ÉTAPE 1/5 : Cahier des Charges Fonctionnel ━━━")
    path_func, functional_data = generate_functional_doc(project_info)
    time.sleep(2)

    # Étape 2 : DAT (basé sur le CdCF)
    logger.info("━━━ ÉTAPE 2/5 : Document d'Architecture Technique ━━━")
    path_tech, technical_data = generate_technical_doc(project_info, functional_data)
    time.sleep(2)

    # Étape 3 : Budget & Planning (basé sur CdCF + DAT)
    logger.info("━━━ ÉTAPE 3/5 : Budget & Planning ━━━")
    path_budget = generate_budget_planning_doc(project_info, functional_data, technical_data)
    time.sleep(1)

    # Étape 4 : Backlog Jira (basé sur le CdCF)
    logger.info("━━━ ÉTAPE 4/5 : Backlog Jira Excel ━━━")
    path_excel = generate_jira_excel(functional_data, project_info["titre"])

    # Étape 5 : Diagramme
    logger.info("━━━ ÉTAPE 5/5 : Diagramme d'Architecture ━━━")
    path_html = generate_html_diagram(project_info)

    elapsed = time.time() - start
    logger.info(f"{'='*60}")
    logger.info(f"GÉNÉRATION TERMINÉE en {elapsed:.0f}s")
    logger.info(f"{'='*60}")

    return {
        "fonctionnel": path_func,
        "technique": path_tech,
        "budget_planning": path_budget,
        "excel": path_excel,
        "html": path_html,
    }
